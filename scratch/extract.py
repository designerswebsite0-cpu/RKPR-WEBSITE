import os
import csv
import openpyxl
from docx import Document
from pypdf import PdfReader

# Paths
rag_root = r"c:\Users\rachi\Downloads\RKPR Website\RKPR_Phase3_RAG_Ready_AUDITED\RKPR_Phase3_RAG_Ready"
output_root = r"c:\Users\rachi\Downloads\RKPR Website\scratch\extracted_rag"

os.makedirs(output_root, exist_ok=True)

def extract_pdf(filepath):
    reader = PdfReader(filepath)
    text = []
    for idx, page in enumerate(reader.pages):
        page_text = page.extract_text()
        text.append(f"--- Page {idx + 1} ---\n{page_text}")
    return "\n\n".join(text)

def extract_docx(filepath):
    doc = Document(filepath)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    
    # Extract tables from docx
    for i, table in enumerate(doc.tables):
        text.append(f"\n--- Table {i + 1} ---")
        for row in table.rows:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            text.append(" | ".join(row_data))
    return "\n".join(text)

def extract_xlsx(filepath):
    wb = openpyxl.load_workbook(filepath, data_only=True)
    text = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text.append(f"--- Sheet: {sheet_name} ---")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                text.append(row_str)
        text.append("\n")
    return "\n".join(text)

def process_file(rel_path):
    full_path = os.path.join(rag_root, rel_path)
    out_dir = os.path.join(output_root, os.path.dirname(rel_path))
    os.makedirs(out_dir, exist_ok=True)
    
    base_name, ext = os.path.splitext(os.path.basename(rel_path))
    out_file = os.path.join(out_dir, f"{base_name}.txt")
    
    print(f"Processing: {rel_path} -> {out_file}")
    
    try:
        if ext.lower() == ".pdf":
            content = extract_pdf(full_path)
        elif ext.lower() == ".docx":
            content = extract_docx(full_path)
        elif ext.lower() in [".xlsx", ".xls"]:
            content = extract_xlsx(full_path)
        elif ext.lower() in [".txt", ".csv", ".json", ".md"]:
            with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        else:
            return
            
        with open(out_file, "w", encoding="utf-8") as f:
            f.write(content)
    except Exception as e:
        print(f"Error processing {rel_path}: {e}")

# Walk through the directories
targets = [
    "00_CONTROL",
    "01_GUEST_KNOWLEDGE",
    "02_MEDIA_INDEXABLE",
    "04_GOVERNANCE",
    "05_STAFF_ONLY"
]

for target in targets:
    target_dir = os.path.join(rag_root, target)
    if not os.path.exists(target_dir):
        continue
    for root, dirs, files in os.walk(target_dir):
        for file in files:
            full_path = os.path.join(root, file)
            rel_path = os.path.relpath(full_path, rag_root)
            process_file(rel_path)

print("Done extracting audited RAG contents.")
